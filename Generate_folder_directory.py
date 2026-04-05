from pathlib import Path
from docx import Document
from datetime import datetime

def export_dir_tree_to_word(folder_path):
    folder = Path(folder_path)
    if not folder.exists():
        print("路径不存在")
        return

    doc = Document()
    doc.add_heading(f'目录结构：{folder.name}', 0)
    doc.add_paragraph(f'路径：{folder.absolute()}')
    doc.add_paragraph(f'时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')

    def add_tree(path, indent=''):
        items = sorted(path.iterdir())
        for i, item in enumerate(items):
            is_last = i == len(items) - 1
            pref = '└── ' if is_last else '├── '

            if item.is_dir():
                doc.add_paragraph(f'{indent}{pref}{item.name}/')
                next_indent = indent + ('    ' if is_last else '│   ')
                add_tree(item, next_indent)
            else:
                doc.add_paragraph(f'{indent}{pref}{item.name}')

    doc.add_paragraph(f'{folder.name}/')
    add_tree(folder)

    # 保存到桌面
    desktop = Path.home() / 'Desktop'
    out = desktop / f'目录结构_{folder.name}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
    doc.save(out)
    print(f'✅ 已保存到桌面：\n{out}')

if __name__ == '__main__':
    export_dir_tree_to_word(r"C:\Users\19868\Desktop\毕业设计\SaleorQA_System")